# utils.py
from docx import Document

def generate_docx_from_minutes(path, title, summary, minutes_struct):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph("Summary:")
    doc.add_paragraph(summary)
    doc.add_heading("Minutes / Action Items", level=2)
    for it in minutes_struct.get("items", []):
        doc.add_paragraph(f"- {it.get('text','')}")
    doc.save(path)
    return path


def transcribe_audio_fallback(path, language="en"):
    """
    Minimal transcription fallback.
    Returns placeholder text if Whisper isn't installed.
    """
    try:
        import whisper
        model = whisper.load_model("small")
        result = model.transcribe(path, language=language)
        return result.get("text", "")
    except Exception:
        return f"[TRANSCRIPTION-PLACEHOLDER] Install whisper for real transcription."
